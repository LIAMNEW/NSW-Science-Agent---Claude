"""
Parse NESA curriculum documents and extract structured content
"""
from docx import Document
import json
import re


def parse_syllabus_support(filepath):
    """Extract teaching advice from NESA syllabus support document"""
    doc = Document(filepath)
    content = {}
    current_focus_area = None
    current_section = None
    section_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Detect focus area headings
        if 'Teaching advice for' in text:
            # Save previous section
            if current_focus_area and current_section:
                if current_focus_area not in content:
                    content[current_focus_area] = {}
                content[current_focus_area][current_section] = '\n'.join(section_content)
            
            # Extract focus area name
            focus_area = text.replace('Teaching advice for', '').strip()
            current_focus_area = focus_area
            section_content = []
            current_section = 'overview'
        
        # Detect sections
        elif current_focus_area:
            if text in ['Key ideas', 'Skills', 'Background knowledge', 'Investigations', 'In context', 'Suggested depth study']:
                # Save previous section
                if current_section and section_content:
                    if current_focus_area not in content:
                        content[current_focus_area] = {}
                    content[current_focus_area][current_section] = '\n'.join(section_content)
                
                current_section = text.lower().replace(' ', '_')
                section_content = []
            elif text:
                section_content.append(text)
    
    # Save last section
    if current_focus_area and current_section and section_content:
        if current_focus_area not in content:
            content[current_focus_area] = {}
        content[current_focus_area][current_section] = '\n'.join(section_content)
    
    return content


def parse_full_syllabus(filepath):
    """Extract outcomes and content from full NESA syllabus"""
    doc = Document(filepath)
    content = {}
    current_focus_area = None
    capturing_content = False
    section_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Detect focus area headings in outcomes section
        if text in ['Observing the Universe', 'Forces', 'Cells and classification', 
                   'Solutions and mixtures', 'Living systems', 'Periodic table and atomic structure', 
                   'Change', 'Data science 1']:
            # Save previous
            if current_focus_area and section_content:
                content[current_focus_area] = '\n'.join(section_content)
            
            current_focus_area = text
            section_content = []
            capturing_content = True
        elif capturing_content and text:
            section_content.append(text)
    
    # Save last
    if current_focus_area and section_content:
        content[current_focus_area] = '\n'.join(section_content)
    
    return content


def extract_stage4_content():
    """Extract and structure all Stage 4 content from NESA documents"""
    
    # Parse syllabus support document
    support_path = "attached_assets/NESA - science_7_10_2023 syllabus support (S4, S5, LS)_1762345248825.docx"
    teaching_advice = parse_syllabus_support(support_path)
    
    # Parse full syllabus
    syllabus_path = "attached_assets/NESA - science_7_10_2023 (S4, S5, LS)_1762345251636.docx"
    outcomes_content = parse_full_syllabus(syllabus_path)
    
    # Combine into structured format
    stage4_content = {}
    
    for focus_area in teaching_advice.keys():
        stage4_content[focus_area] = {
            'teaching_advice': teaching_advice.get(focus_area, {}),
            'outcomes': outcomes_content.get(focus_area, '')
        }
    
    return stage4_content


if __name__ == '__main__':
    # Test extraction
    content = extract_stage4_content()
    
    print("Extracted NESA Content:")
    print("=" * 60)
    for area, data in content.items():
        print(f"\n{area}:")
        print(f"  - Teaching advice sections: {list(data['teaching_advice'].keys())}")
        print(f"  - Outcomes length: {len(data['outcomes'])} characters")
